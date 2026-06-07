# -*- coding: utf-8 -*-
"""
根据 report_a_data.json / report_b_data.json 生成两份 docx 报告。
脚本本身纯 ASCII，所有中文通过 JSON 数据文件传入，避免编码损坏。
"""
import json
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = r"f:\Claude Project\大数据应用系统开发实践\agent5_report"
FONT_HEITI = "黑体"
FONT_SONGTI = "宋体"
FONT_EN = "Times New Roman"
SIZE_H1 = Pt(16)
SIZE_H2 = Pt(12)
SIZE_BODY = Pt(12)
LINE_SPACING_PT = 22


def set_cn_font(run, font_name):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), FONT_EN)
    rFonts.set(qn('w:hAnsi'), FONT_EN)


def set_run_font(run, size=None, bold=None, cn_font=None):
    if size: run.font.size = size
    if bold is not None: run.font.bold = bold
    run.font.name = FONT_EN
    if cn_font: set_cn_font(run, cn_font)


def set_para_fmt(para, alignment=None, indent=None, sb=None, sa=None, line_pt=LINE_SPACING_PT):
    pf = para.paragraph_format
    if alignment is not None: para.alignment = alignment
    if indent is not None: pf.first_line_indent = indent
    if sb is not None: pf.space_before = sb
    if sa is not None: pf.space_after = sa
    if line_pt:
        pf.line_spacing = Pt(line_pt)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY


def add_h1(doc, text):
    p = doc.add_paragraph()
    set_para_fmt(p, WD_ALIGN_PARAGRAPH.CENTER, sb=Pt(12), sa=Pt(12))
    r = p.add_run(text)
    set_run_font(r, size=SIZE_H1, bold=True, cn_font=FONT_HEITI)


def add_h2(doc, text):
    p = doc.add_paragraph()
    set_para_fmt(p, WD_ALIGN_PARAGRAPH.LEFT, sb=Pt(6), sa=Pt(6))
    r = p.add_run(text)
    set_run_font(r, size=SIZE_H2, bold=True, cn_font=FONT_HEITI)


def add_h3(doc, text):
    p = doc.add_paragraph()
    set_para_fmt(p, WD_ALIGN_PARAGRAPH.LEFT, sb=Pt(3), sa=Pt(3))
    r = p.add_run(text)
    set_run_font(r, size=SIZE_H2, bold=False, cn_font=FONT_HEITI)


def add_p(doc, text):
    p = doc.add_paragraph()
    set_para_fmt(p, WD_ALIGN_PARAGRAPH.JUSTIFY, indent=Pt(SIZE_BODY.pt * 2))
    r = p.add_run(text)
    set_run_font(r, size=SIZE_BODY, cn_font=FONT_SONGTI)


def add_code(doc, code):
    for line in code.split('\n'):
        p = doc.add_paragraph()
        set_para_fmt(p, WD_ALIGN_PARAGRAPH.LEFT, indent=Pt(0),
                     sb=Pt(0), sa=Pt(0), line_pt=18)
        r = p.add_run(line if line else ' ')
        r.font.name = "Consolas"
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:ascii'), "Consolas")
        rFonts.set(qn('w:hAnsi'), "Consolas")
        rFonts.set(qn('w:eastAsia'), "Consolas")
        r.font.size = Pt(10)


def add_table(doc, grid):
    if not grid: return
    rows, cols = len(grid), len(grid[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(grid):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ''
            p = cell.paragraphs[0]
            set_para_fmt(p, WD_ALIGN_PARAGRAPH.CENTER, indent=Pt(0),
                         sb=Pt(0), sa=Pt(0))
            r = p.add_run(str(val))
            set_run_font(r, size=SIZE_BODY, bold=(i == 0), cn_font=FONT_SONGTI)


def add_pagebreak(doc):
    doc.add_page_break()


def add_footer_pagenum(doc):
    section = doc.sections[0]
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_fmt(p, WD_ALIGN_PARAGRAPH.CENTER, line_pt=18)
    r1 = p.add_run("第 - ")
    r2 = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = " PAGE "
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    r2._element.append(fldChar1)
    r2._element.append(instrText)
    r2._element.append(fldChar2)
    r3 = p.add_run(" - 页")
    for r in (r1, r2, r3):
        set_run_font(r, size=Pt(10), cn_font=FONT_SONGTI)


def apply_default_style(doc):
    style = doc.styles['Normal']
    style.font.name = FONT_EN
    style.font.size = SIZE_BODY
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_SONGTI)
    rFonts.set(qn('w:ascii'), FONT_EN)
    rFonts.set(qn('w:hAnsi'), FONT_EN)
    pf = style.paragraph_format
    pf.line_spacing = Pt(LINE_SPACING_PT)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.17)
        sec.right_margin = Cm(3.17)


def build_cover(doc, title, name, sid, klass):
    for _ in range(2):
        p = doc.add_paragraph()
        set_para_fmt(p)

    p = doc.add_paragraph()
    set_para_fmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("重庆交通大学信息科学与工程学院")
    set_run_font(r, size=Pt(18), bold=True, cn_font=FONT_HEITI)

    p = doc.add_paragraph()
    set_para_fmt(p, WD_ALIGN_PARAGRAPH.CENTER, sb=Pt(6))
    r = p.add_run("课  程  设  计  报  告")
    set_run_font(r, size=Pt(22), bold=True, cn_font=FONT_HEITI)

    for _ in range(3):
        p = doc.add_paragraph()
        set_para_fmt(p)

    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in table.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(11.5)
    info = [
        ("题    目", title),
        ("课程名称", "大数据应用系统开发实践"),
        ("专业班级", klass),
        ("学    号", sid),
        ("姓    名", name),
        ("任课教师", "李  韧"),
    ]
    for i, (k, v) in enumerate(info):
        c0 = table.cell(i, 0)
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c0.text = ''
        p0 = c0.paragraphs[0]
        set_para_fmt(p0, WD_ALIGN_PARAGRAPH.CENTER, indent=Pt(0))
        r0 = p0.add_run(k)
        set_run_font(r0, size=SIZE_BODY, bold=True, cn_font=FONT_HEITI)

        c1 = table.cell(i, 1)
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c1.text = ''
        p1 = c1.paragraphs[0]
        set_para_fmt(p1, WD_ALIGN_PARAGRAPH.LEFT, indent=Pt(0))
        r1 = p1.add_run(v)
        set_run_font(r1, size=SIZE_BODY, cn_font=FONT_SONGTI)

    for _ in range(4):
        p = doc.add_paragraph()
        set_para_fmt(p)

    p = doc.add_paragraph()
    set_para_fmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("2026 年 6 月")
    set_run_font(r, size=Pt(16), bold=True, cn_font=FONT_HEITI)

    add_pagebreak(doc)


def build_toc(doc, sections):
    p = doc.add_paragraph()
    set_para_fmt(p, WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("目  录")
    set_run_font(r, size=SIZE_H1, bold=True, cn_font=FONT_HEITI)
    doc.add_paragraph()
    for title, page in sections:
        p = doc.add_paragraph()
        set_para_fmt(p, WD_ALIGN_PARAGRAPH.LEFT, indent=Pt(0))
        # 用点号填充使页码右对齐
        gap = max(2, 60 - len(title))
        r = p.add_run(f"{title}{'.' * gap}{page}")
        set_run_font(r, size=SIZE_BODY, cn_font=FONT_SONGTI)
    add_pagebreak(doc)


def render_body(doc, body):
    for item in body:
        t = item['type']
        if t == 'h1': add_h1(doc, item['text'])
        elif t == 'h2': add_h2(doc, item['text'])
        elif t == 'h3': add_h3(doc, item['text'])
        elif t == 'p': add_p(doc, item['text'])
        elif t == 'code': add_code(doc, item['text'])
        elif t == 'table': add_table(doc, item['rows'])
        elif t == 'pagebreak': add_pagebreak(doc)


def build_one(data_path, out_name):
    with open(data_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    doc = Document()
    apply_default_style(doc)
    add_footer_pagenum(doc)
    build_cover(doc, data['title'], data['author'], data['student_id'], data['klass'])
    build_toc(doc, data['toc'])
    render_body(doc, data['body'])
    out = os.path.join(OUT_DIR, out_name)
    doc.save(out)
    print(f"[OK] {out}  ({os.path.getsize(out)/1024:.1f} KB)")
    return out


if __name__ == "__main__":
    a = build_one(os.path.join(OUT_DIR, "report_a_data.json"),
                  "报告A_数据层_课程设计报告_v2.docx")
    b = build_one(os.path.join(OUT_DIR, "report_b_data.json"),
                  "报告B_应用层_课程设计报告_v2.docx")
